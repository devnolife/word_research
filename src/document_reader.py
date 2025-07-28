"""
Core document reader module for Word Research Analyzer
"""
import os
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
from dataclasses import dataclass, field
from datetime import datetime
import docx
from docx.document import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph

from .config import DEFAULT_DOCUMENT_CONFIG, DEFAULT_PROCESSING_CONFIG
from .utils import (
    validate_file_path, get_file_size, get_file_extension, 
    is_supported_format, calculate_file_hash, clean_text, format_bytes
)
from .exceptions import (
    DocumentProcessingError, FileValidationError, DocumentParsingError,
    UnsupportedFormatError, FileSizeError, CorruptDocumentError, MemoryError
)
from .logging_config import get_logger

logger = get_logger(__name__)

@dataclass
class DocumentStats:
    """Document statistics data structure"""
    total_paragraphs: int = 0
    total_words: int = 0
    total_characters: int = 0
    total_characters_no_spaces: int = 0
    total_sentences: int = 0
    total_pages: int = 0
    total_sections: int = 0
    total_tables: int = 0
    total_images: int = 0
    total_headings: int = 0
    reading_time_minutes: float = 0.0
    complexity_score: float = 0.0
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return {
            'total_paragraphs': self.total_paragraphs,
            'total_words': self.total_words,
            'total_characters': self.total_characters,
            'total_characters_no_spaces': self.total_characters_no_spaces,
            'total_sentences': self.total_sentences,
            'total_pages': self.total_pages,
            'total_sections': self.total_sections,
            'total_tables': self.total_tables,
            'total_images': self.total_images,
            'total_headings': self.total_headings,
            'reading_time_minutes': self.reading_time_minutes,
            'complexity_score': self.complexity_score
        }

@dataclass
class ParagraphInfo:
    """Paragraph information data structure"""
    index: int
    text: str
    style_name: str
    is_heading: bool = False
    heading_level: Optional[int] = None
    alignment: Optional[str] = None
    bold: bool = False
    italic: bool = False
    underline: bool = False
    font_size: Optional[float] = None
    font_name: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return {
            'index': self.index,
            'text': self.text,
            'style_name': self.style_name,
            'is_heading': self.is_heading,
            'heading_level': self.heading_level,
            'alignment': self.alignment,
            'bold': self.bold,
            'italic': self.italic,
            'underline': self.underline,
            'font_size': self.font_size,
            'font_name': self.font_name
        }

@dataclass
class HeadingInfo:
    """Heading information data structure"""
    index: int
    text: str
    level: int
    style_name: str
    paragraph_index: int
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return {
            'index': self.index,
            'text': self.text,
            'level': self.level,
            'style_name': self.style_name,
            'paragraph_index': self.paragraph_index
        }

@dataclass
class TableInfo:
    """Table information data structure"""
    index: int
    rows: int
    columns: int
    data: List[List[str]] = field(default_factory=list)
    style_name: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return {
            'index': self.index,
            'rows': self.rows,
            'columns': self.columns,
            'data': self.data,
            'style_name': self.style_name
        }

class DocumentReader:
    """
    Core document reader class for analyzing Word documents
    """
    
    def __init__(self, config=None, processing_config=None):
        """
        Initialize DocumentReader
        
        Args:
            config: Document configuration
            processing_config: Processing configuration
        """
        self.config = config or DEFAULT_DOCUMENT_CONFIG
        self.processing_config = processing_config or DEFAULT_PROCESSING_CONFIG
        
        self.document: Optional[DocxDocument] = None
        self.file_path: Optional[Path] = None
        self.file_info: Dict[str, Any] = {}
        self.is_loaded: bool = False
        
    def load_document(self, file_path: Union[str, Path]) -> bool:
        """
        Load Word document from file path
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            bool: True if document loaded successfully
            
        Raises:
            FileValidationError: If file validation fails
            DocumentParsingError: If document parsing fails
            CorruptDocumentError: If document is corrupted
        """
        try:
            logger.info(f"Starting document load process for: {file_path}")
            
            # Validate file
            if not self.validate_file(file_path):
                error_msg = f"File validation failed for: {file_path}"
                logger.error(error_msg)
                raise FileValidationError(str(file_path), "File validation failed")
            
            # Store file information
            self.file_path = Path(file_path)
            file_size = get_file_size(file_path)
            
            self.file_info = {
                'path': str(self.file_path),
                'name': self.file_path.name,
                'size_bytes': file_size,
                'size_formatted': format_bytes(file_size),
                'extension': get_file_extension(file_path),
                'hash': calculate_file_hash(file_path),
                'loaded_at': datetime.now().isoformat()
            }
            
            logger.debug(f"File info collected: {self.file_info}")
            
            # Load document
            try:
                self.document = docx.Document(file_path)
                self.is_loaded = True
                logger.info(f"Document loaded successfully: {file_path}")
                return True
                
            except docx.opc.exceptions.PackageNotFoundError as e:
                logger.error(f"Document file not found or corrupted: {e}")
                raise CorruptDocumentError(str(file_path), e)
                
            except docx.oxml.exceptions.InvalidXmlError as e:
                logger.error(f"Invalid XML structure in document: {e}")
                raise CorruptDocumentError(str(file_path), e)
                
            except PermissionError as e:
                logger.error(f"Permission denied accessing file: {e}")
                raise FileValidationError(str(file_path), f"Permission denied: {e}")
                
            except MemoryError as e:
                logger.error(f"Memory error loading document: {e}")
                raise MemoryError(str(file_path), "document_loading")
                
        except (FileValidationError, DocumentParsingError, CorruptDocumentError, MemoryError):
            # Re-raise our custom exceptions
            self.is_loaded = False
            raise
            
        except Exception as e:
            # Catch all other exceptions and wrap them
            self.is_loaded = False
            logger.error(f"Unexpected error loading document: {e}")
            raise DocumentParsingError(str(file_path), "document_loading", e)
    
    def validate_file(self, file_path: Union[str, Path]) -> bool:
        """
        Validate file format and existence
        
        Args:
            file_path: Path to the file
            
        Returns:
            bool: True if file is valid
            
        Raises:
            FileValidationError: If file validation fails
            UnsupportedFormatError: If file format is not supported
            FileSizeError: If file size exceeds limits
        """
        logger.debug(f"Validating file: {file_path}")
        
        # Check if file exists and is readable
        if not validate_file_path(file_path):
            error_msg = f"File does not exist or is not readable: {file_path}"
            logger.error(error_msg)
            raise FileValidationError(str(file_path), "File does not exist or is not readable")
        
        # Check file extension
        extension = get_file_extension(file_path)
        if not is_supported_format(file_path, self.config.supported_extensions):
            logger.error(f"Unsupported file format: {extension}")
            raise UnsupportedFormatError(str(file_path), extension, self.config.supported_extensions)
        
        # Check file size
        file_size_bytes = get_file_size(file_path)
        file_size_mb = file_size_bytes / (1024 * 1024)
        
        if file_size_mb > self.config.max_file_size_mb:
            logger.error(f"File size {file_size_mb:.2f}MB exceeds limit {self.config.max_file_size_mb}MB")
            raise FileSizeError(str(file_path), file_size_mb, self.config.max_file_size_mb)
        
        logger.debug(f"File validation passed for: {file_path}")
        return True
    
    def get_basic_info(self) -> Dict[str, Any]:
        """
        Get basic document information
        
        Returns:
            Dict[str, Any]: Basic document statistics
        """
        if not self.is_loaded:
            return {}
        
        stats = self.get_document_statistics()
        
        return {
            'file_info': self.file_info,
            'paragraphs': stats.total_paragraphs,
            'words': stats.total_words,
            'characters': stats.total_characters,
            'characters_no_spaces': stats.total_characters_no_spaces,
            'tables': stats.total_tables,
            'headings': stats.total_headings,
            'is_loaded': self.is_loaded
        }
    
    def extract_raw_text(self) -> str:
        """
        Extract all raw text content from document
        
        Returns:
            str: Raw text content
        """
        if not self.is_loaded:
            return ""
        
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables if enabled
        if self.processing_config.extract_tables:
            for table in self.document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
        
        raw_text = "\n".join(text_parts)
        return clean_text(raw_text, preserve_whitespace=True)
    
    def get_document_metadata(self) -> Dict[str, Any]:
        """
        Extract document metadata
        
        Returns:
            Dict[str, Any]: Document metadata
        """
        if not self.is_loaded:
            return {}
        
        core_props = self.document.core_properties
        
        metadata = {
            'title': core_props.title or "",
            'author': core_props.author or "",
            'subject': core_props.subject or "",
            'keywords': core_props.keywords or "",
            'comments': core_props.comments or "",
            'category': core_props.category or "",
            'created': core_props.created.isoformat() if core_props.created else None,
            'modified': core_props.modified.isoformat() if core_props.modified else None,
            'last_modified_by': core_props.last_modified_by or "",
            'revision': core_props.revision,
            'version': core_props.version or "",
            'language': core_props.language or ""
        }
        
        return metadata
    
    def extract_paragraphs(self) -> List[ParagraphInfo]:
        """
        Extract paragraphs with structure information
        
        Returns:
            List[ParagraphInfo]: List of paragraph information
        """
        if not self.is_loaded:
            return []
        
        paragraphs = []
        
        for i, paragraph in enumerate(self.document.paragraphs):
            if not paragraph.text.strip():
                continue
                
            # Get style information
            style_name = paragraph.style.name if paragraph.style else "Normal"
            
            # Check if it's a heading
            is_heading = style_name.startswith('Heading') or style_name.startswith('Title')
            heading_level = None
            
            if is_heading and style_name.startswith('Heading'):
                try:
                    heading_level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    heading_level = 1
            
            # Get formatting information
            runs = paragraph.runs
            bold = any(run.bold for run in runs if run.bold is not None)
            italic = any(run.italic for run in runs if run.italic is not None)
            underline = any(run.underline for run in runs if run.underline is not None)
            
            # Get font information from first run
            font_size = None
            font_name = None
            if runs:
                first_run = runs[0]
                if first_run.font.size:
                    font_size = first_run.font.size.pt
                font_name = first_run.font.name
            
            # Get alignment
            alignment = str(paragraph.alignment) if paragraph.alignment else None
            
            paragraph_info = ParagraphInfo(
                index=i,
                text=clean_text(paragraph.text),
                style_name=style_name,
                is_heading=is_heading,
                heading_level=heading_level,
                alignment=alignment,
                bold=bold,
                italic=italic,
                underline=underline,
                font_size=font_size,
                font_name=font_name
            )
            
            paragraphs.append(paragraph_info)
        
        return paragraphs
    
    def identify_headings(self) -> List[HeadingInfo]:
        """
        Identify headings and their levels
        
        Returns:
            List[HeadingInfo]: List of heading information
        """
        if not self.is_loaded:
            return []
        
        headings = []
        heading_index = 0
        
        for para_index, paragraph in enumerate(self.document.paragraphs):
            if not paragraph.text.strip():
                continue
                
            style_name = paragraph.style.name if paragraph.style else "Normal"
            
            # Check if it's a heading
            if style_name.startswith('Heading') or style_name.startswith('Title'):
                level = 1  # Default level
                
                if style_name.startswith('Heading'):
                    try:
                        level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                elif style_name.startswith('Title'):
                    level = 0  # Title is higher than Heading 1
                
                heading_info = HeadingInfo(
                    index=heading_index,
                    text=clean_text(paragraph.text),
                    level=level,
                    style_name=style_name,
                    paragraph_index=para_index
                )
                
                headings.append(heading_info)
                heading_index += 1
        
        return headings
    
    def extract_tables_basic(self) -> List[TableInfo]:
        """
        Extract basic table information
        
        Returns:
            List[TableInfo]: List of table information
        """
        if not self.is_loaded:
            return []
        
        tables = []
        
        for i, table in enumerate(self.document.tables):
            rows = len(table.rows)
            columns = len(table.columns) if table.rows else 0
            
            # Extract table data
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = clean_text(cell.text)
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            # Get table style if available
            style_name = None
            try:
                if hasattr(table, 'style') and table.style:
                    style_name = table.style.name
            except:
                pass
            
            table_info = TableInfo(
                index=i,
                rows=rows,
                columns=columns,
                data=table_data,
                style_name=style_name
            )
            
            tables.append(table_info)
        
        return tables
    
    def get_document_statistics(self) -> DocumentStats:
        """
        Get comprehensive document statistics
        
        Returns:
            DocumentStats: Document statistics
        """
        if not self.is_loaded:
            logger.warning("Document not loaded, returning empty statistics")
            return DocumentStats()
        
        logger.debug("Calculating document statistics...")
        stats = DocumentStats()
        
        # Count paragraphs and analyze text
        total_text = ""
        paragraph_count = 0
        heading_count = 0
        
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                paragraph_count += 1
                total_text += paragraph.text + " "
                
                # Check if it's a heading
                style_name = paragraph.style.name if paragraph.style else "Normal"
                if style_name.startswith('Heading') or style_name.startswith('Title'):
                    heading_count += 1
        
        stats.total_paragraphs = paragraph_count
        stats.total_headings = heading_count
        
        # Count words and characters
        if total_text:
            words = total_text.split()
            stats.total_words = len(words)
            stats.total_characters = len(total_text)
            stats.total_characters_no_spaces = len(total_text.replace(' ', ''))
            
            # More sophisticated sentence counting
            import re
            sentences = re.split(r'[.!?]+', total_text)
            # Filter out empty sentences and very short ones (likely not real sentences)
            sentences = [s.strip() for s in sentences if len(s.strip()) > 10]
            stats.total_sentences = len(sentences)
            
            # Calculate reading time (average 200 words per minute)
            stats.reading_time_minutes = round(stats.total_words / 200.0, 2)
            
            # Calculate complexity score (basic implementation)
            stats.complexity_score = self._calculate_complexity_score(
                stats.total_words, 
                stats.total_sentences, 
                stats.total_paragraphs,
                heading_count
            )
        
        # Count tables
        stats.total_tables = len(self.document.tables)
        
        # Count sections
        stats.total_sections = len(self.document.sections)
        
        # Estimate page count (more sophisticated than before)
        words_per_page = 250  # Standard assumption
        if stats.total_words > 0:
            stats.total_pages = max(1, round(stats.total_words / words_per_page))
        else:
            stats.total_pages = 1
        
        # Note: Image count requires more complex processing - placeholder for now
        stats.total_images = 0  # Will implement in future phases
        
        logger.debug(f"Statistics calculated: {stats.total_words} words, {stats.total_paragraphs} paragraphs")
        
        return stats
    
    def _calculate_complexity_score(self, words: int, sentences: int, paragraphs: int, headings: int) -> float:
        """
        Calculate document complexity score (0-10 scale)
        
        Args:
            words: Total word count
            sentences: Total sentence count 
            paragraphs: Total paragraph count
            headings: Total heading count
            
        Returns:
            float: Complexity score between 0.0 and 10.0
        """
        if words == 0 or sentences == 0:
            return 0.0
        
        # Average words per sentence (higher = more complex)
        avg_words_per_sentence = words / sentences
        
        # Average sentences per paragraph (higher = more complex)
        avg_sentences_per_paragraph = sentences / max(paragraphs, 1)
        
        # Document length factor
        length_factor = min(words / 1000, 3.0)  # Cap at 3.0 for very long docs
        
        # Structure complexity (more headings = better structure = lower complexity)
        structure_factor = max(0.5, 1.0 - (headings / max(paragraphs, 1)))
        
        # Combine factors into complexity score
        complexity = (
            (avg_words_per_sentence / 20.0) * 3.0 +  # Max 3 points for sentence length
            (avg_sentences_per_paragraph / 5.0) * 2.0 +  # Max 2 points for paragraph density
            length_factor * 2.0 +  # Max 2 points for document length
            structure_factor * 3.0  # Max 3 points for poor structure
        )
        
        # Normalize to 0-10 scale and round
        return round(min(complexity, 10.0), 2)
